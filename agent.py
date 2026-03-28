"""
ApplyWiz Engineering Challenge — Option 2: AI Resume Tailoring Agent

Pipeline:
  1. Parse option2_job_links.xlsx and option2_jobs.json, join on id/#
  2. Load candidate resume from resume/candidate_resume.docx
  3. For each job, call Groq AI (Llama 3.3-70b) to tailor the resume
  4. Save tailored resume as a Word document (.docx)
  5. Email the tailored resume with job details in the body
  6. Each job is processed independently — one failure never stops the rest
"""

import os
import json
import time
import smtplib
import traceback
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

import pandas as pd
from dotenv import load_dotenv
from groq import Groq
from docx import Document
from docx.shared import Inches, Pt

load_dotenv()

# ── Configuration ─────────────────────────────────────────────────────────────
GROQ_API_KEY    = os.getenv("GROQ_API_KEY")
GMAIL_USER      = os.getenv("GMAIL_USER")
GMAIL_PASSWORD  = os.getenv("GMAIL_APP_PASSWORD")
RECIPIENT_EMAIL = os.getenv("RECIPIENT_EMAIL")

GROQ_MODEL  = "llama-3.3-70b-versatile"
OUTPUT_DIR  = "tailored_resumes"
MAX_RETRIES = 3  # Retry AI calls up to 3 times on failure

os.makedirs(OUTPUT_DIR, exist_ok=True)
client = Groq(api_key=GROQ_API_KEY)


# ── Data Loading ──────────────────────────────────────────────────────────────

def load_jobs() -> list[dict]:
    """
    Load job metadata from Excel and full descriptions from JSON.
    Join the two sources on the id / # field.
    """
    df = pd.read_excel("option2_job_links.xlsx").dropna(subset=["#"])
    df["#"] = df["#"].astype(int)

    with open("option2_jobs.json", "r") as f:
        jobs_json = json.load(f)["jobs"]

    jobs_map = {j["id"]: j for j in jobs_json}

    jobs = []
    for _, row in df.iterrows():
        job_id = int(row["#"])
        if job_id in jobs_map:
            detail = jobs_map[job_id]
            jobs.append({
                "id":           job_id,
                "title":        row["Job Title"],
                "company":      row["Company"],
                "url":          row["URL"],
                "description":  detail.get("description", ""),
                "requirements": detail.get("requirements", []),
                "nice_to_have": detail.get("nice_to_have", []),
            })

    return jobs


def read_resume(path: str) -> str:
    """Extract plain text from the candidate's .docx resume."""
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


# ── AI Tailoring ──────────────────────────────────────────────────────────────

def build_prompt(resume_text: str, job: dict) -> str:
    """Build the prompt sent to the LLM for resume tailoring."""
    requirements = "\n".join(f"- {r}" for r in job["requirements"])
    nice_to_have  = "\n".join(f"- {n}" for n in job["nice_to_have"])

    return f"""You are an expert technical resume writer.

Rewrite the candidate resume below so it is strongly tailored to this specific job.

RULES:
1. Keep ALL factual information (names, companies, dates, degrees) exactly as-is — do NOT invent experience.
2. Reorder bullet points so the most relevant ones appear first for this role.
3. Reword bullets to mirror the language and keywords in the job description naturally.
4. Rewrite the professional summary to speak directly to this role.
5. Add a KEY SKILLS section at the top listing skills most relevant to this job.
6. The output should look noticeably different from the original — not just minor rewording.
7. Use only plain ASCII characters — use hyphen (-) instead of bullet symbols or special dashes.
8. Return ONLY the full resume text with section headers in ALL CAPS. No extra commentary.

JOB TITLE: {job['title']}
COMPANY: {job['company']}

DESCRIPTION:
{job['description']}

REQUIREMENTS:
{requirements}

NICE TO HAVE:
{nice_to_have}

ORIGINAL RESUME:
{resume_text}

Write the tailored resume now:"""


def tailor_resume(resume_text: str, job: dict) -> str:
    """
    Call Groq AI to tailor the resume for a specific job.
    Retries up to MAX_RETRIES times with exponential backoff on failure.
    """
    prompt = build_prompt(resume_text, job)

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            response = client.chat.completions.create(
                model=GROQ_MODEL,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.4,
            )
            return response.choices[0].message.content.strip()

        except Exception as e:
            if attempt < MAX_RETRIES:
                wait = 2 ** attempt  # 2s, 4s, 8s
                print(f"      -> Attempt {attempt} failed, retrying in {wait}s...")
                time.sleep(wait)
            else:
                raise e  # All retries exhausted — let caller handle it


# ── Document Generation ───────────────────────────────────────────────────────

def save_as_docx(text: str, filepath: str):
    """
    Save the tailored resume text as a formatted Word document.
    Section headers (ALL CAPS lines) are rendered bold and slightly larger.
    """
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin    = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin   = Inches(0.9)
    section.right_margin  = Inches(0.9)

    for line in text.split("\n"):
        line = line.strip()

        if not line:
            doc.add_paragraph("")
            continue

        # Detect section headers: ALL CAPS or short line ending with colon
        is_header = line.isupper() or (line.endswith(":") and len(line) < 40)

        p   = doc.add_paragraph()
        run = p.add_run(line)
        if is_header:
            run.bold      = True
            run.font.size = Pt(12)
        else:
            run.font.size = Pt(10)

    doc.save(filepath)


# ── Email Delivery ─────────────────────────────────────────────────────────────

def send_email(job: dict, file_path: str):
    """
    Send one email with the tailored resume attached.
    Includes job title, company, and URL in the email body.
    """
    msg             = MIMEMultipart()
    msg["From"]     = GMAIL_USER
    msg["To"]       = RECIPIENT_EMAIL
    msg["Subject"]  = f"Tailored Resume - {job['title']} at {job['company']}"

    body = f"""Hi,

Please find attached a resume tailored specifically for the following role:

  Job Title : {job['title']}
  Company   : {job['company']}
  Job URL   : {job['url']}

This resume has been customised by an AI agent to emphasise the skills and
experience most relevant to this specific position.

Best regards,
ApplyWiz Resume Agent
"""
    msg.attach(MIMEText(body, "plain"))

    # Attach the Word document
    with open(file_path, "rb") as f:
        part = MIMEBase("application", "octet-stream")
        part.set_payload(f.read())
    encoders.encode_base64(part)
    part.add_header(
        "Content-Disposition",
        f'attachment; filename="{os.path.basename(file_path)}"',
    )
    msg.attach(part)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
        server.login(GMAIL_USER, GMAIL_PASSWORD)
        server.sendmail(GMAIL_USER, RECIPIENT_EMAIL, msg.as_string())


# ── Main Pipeline ─────────────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("  ApplyWiz - AI Resume Tailoring Agent")
    print(f"  Model : {GROQ_MODEL}")
    print("=" * 60)

    # Load all inputs
    jobs = load_jobs()
    print(f"\n[OK] Loaded {len(jobs)} jobs from Excel + JSON")

    resume_text = read_resume("resume/candidate_resume.docx")
    print(f"[OK] Resume loaded ({len(resume_text.split())} words)\n")

    results = []

    for job in jobs:
        print(f"[{job['id']}/5] {job['title']} - {job['company']}")
        status = {"job": job["title"], "status": "success", "error": None}

        try:
            # Step 1: Tailor resume via AI
            print(f"      -> Tailoring resume with Groq AI...")
            tailored = tailor_resume(resume_text, job)

            # Step 2: Save as Word document
            safe_title = job["title"].lower().replace(" ", "_").replace("/", "_")
            file_name  = f"resume_tailored_{safe_title}.docx"
            file_path  = os.path.join(OUTPUT_DIR, file_name)
            save_as_docx(tailored, file_path)
            print(f"      -> Saved : {file_path}")

            # Step 3: Send email with attachment
            print(f"      -> Emailing to {RECIPIENT_EMAIL}...")
            send_email(job, file_path)
            print(f"      -> [OK] Done!\n")

        except Exception as e:
            # Log failure but continue processing remaining jobs
            status["status"] = "failed"
            status["error"]  = str(e)
            print(f"      -> [FAILED] {e}\n")
            traceback.print_exc()

        results.append(status)

    # ── Summary ───────────────────────────────────────────────────────────────
    success = sum(1 for r in results if r["status"] == "success")
    failed  = len(results) - success

    print("=" * 60)
    print(f"  PIPELINE COMPLETE — {success} succeeded, {failed} failed")
    print("=" * 60)
    for r in results:
        icon = "OK    " if r["status"] == "success" else "FAILED"
        print(f"  [{icon}]  {r['job']}")
        if r["error"]:
            print(f"            Error: {r['error']}")
    print("=" * 60)


if __name__ == "__main__":
    main()